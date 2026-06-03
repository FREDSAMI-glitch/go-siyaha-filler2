from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, JSONResponse
from docxtpl import DocxTemplate
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries, column_index_from_string
from pathlib import Path
from jinja2 import Environment, ChainableUndefined
from docx import Document
from docx.table import _Cell
from docx.shared import Cm
from docx.oxml.ns import qn
from lxml import etree
from datetime import date
import tempfile
import zipfile
import uuid
import json
import re
import copy

app = FastAPI()

CODE_VERSION = "DAP_CLEAN_TABLES_ADDRESS_V6_2026_06_03"

API_DIR = Path(__file__).resolve().parent
POSSIBLE_BASE_DIRS = [API_DIR, API_DIR.parent]

BASE_DIR = None
for d in POSSIBLE_BASE_DIRS:
    if (d / "templates").exists() or (d / "mappings").exists():
        BASE_DIR = d
        break

if BASE_DIR is None:
    BASE_DIR = API_DIR.parent

TEMPLATE_DIR = BASE_DIR / "templates"
MAPPING_DIR = BASE_DIR / "mappings"

DOCX_JINJA_ENV = Environment(undefined=ChainableUndefined)


# =========================================================
# OUTILS GÉNÉRAUX
# =========================================================

def safe_name(value):
    if not value:
        return "DOSSIER"
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in str(value))


def normalize_key_text(value):
    text = str(value or "").lower()
    text = text.replace("é", "e").replace("è", "e").replace("ê", "e")
    text = text.replace("à", "a").replace("â", "a")
    text = text.replace("ç", "c").replace("ï", "i").replace("î", "i")
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


def as_dict(value):
    return value if isinstance(value, dict) else {}


def is_empty(value):
    return value is None or value == "" or value == [] or value == {}


def deep_get(data, path, default=None):
    if not path:
        return default

    current = data

    for part in str(path).split("."):
        if isinstance(current, dict):
            if part in current:
                current = current.get(part)
            else:
                return default
        elif isinstance(current, list):
            try:
                current = current[int(part)]
            except Exception:
                return default
        else:
            return default

    return current if current is not None else default


def deep_set(data, path, value):
    parts = str(path).split(".")
    current = data

    for part in parts[:-1]:
        if not isinstance(current.get(part), dict):
            current[part] = {}
        current = current[part]

    current[parts[-1]] = value


def get_context_value(context, path, default=""):
    value = deep_get(context, path, None)

    if value not in [None, "", [], {}]:
        return value

    dap = context.get("dap", {})
    if isinstance(dap, dict):
        value = deep_get(dap, path, None)
        if value not in [None, "", [], {}]:
            return value

    return default


def first_non_empty(context, paths, default=""):
    for path in paths:
        value = get_context_value(context, path, None)
        if value not in [None, "", [], {}]:
            return value
    return default


def parse_number(value):
    if value is None or value == "":
        return None

    if isinstance(value, bool):
        return int(value)

    if isinstance(value, (int, float)):
        return value

    text = str(value).strip()
    text = text.replace("\u00a0", " ")

    if text.endswith("%"):
        try:
            return float(text.replace("%", "").replace(",", ".")) / 100
        except Exception:
            return None

    text = re.sub(r"[^\d,.\-]", "", text)

    if not text:
        return None

    if text.count(",") > 1 and "." not in text:
        text = text.replace(",", "")
    elif text.count(".") > 1 and "," not in text:
        text = text.replace(".", "")
    elif "," in text and "." in text:
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    elif "," in text and "." not in text:
        text = text.replace(",", ".")

    try:
        number = float(text)
        if number.is_integer():
            return int(number)
        return number
    except Exception:
        return None


def to_number(value):
    number = parse_number(value)
    return 0 if number is None else number


def number_as_kmad(value, source_path=""):
    number = to_number(value)
    if not number:
        return 0
    if str(source_path).endswith("_kmad"):
        return number
    if abs(number) >= 100000:
        return number / 1000
    return number


def number_as_mad(value, source_path=""):
    number = to_number(value)
    if not number:
        return 0
    if str(source_path).endswith("_kmad"):
        return number * 1000
    return number


def format_display_value(value, fmt=None):
    if value is None:
        return ""

    if fmt in ["mad", "money"]:
        number = to_number(value)
        if number == 0:
            return ""
        return f"{number:,.0f}".replace(",", " ") + " MAD"

    if fmt == "kmad":
        number = to_number(value)
        if number == 0:
            return ""
        return f"{number:,.0f}".replace(",", " ") + " KMAD"

    if fmt == "mmad_an":
        number = to_number(value)
        if number == 0:
            return ""
        return f"{number:,.1f}".replace(",", " ") + " MMAD/an"

    if fmt == "percent":
        number = to_number(value)
        if number == 0:
            return ""
        if number <= 1:
            number = number * 100
        return f"{number:.0f}%"

    if fmt in ["number", "integer", "float"]:
        number = to_number(value)
        return str(int(number)) if number == int(number) else str(number)

    return str(value)


def find_file(filename):
    possible_paths = [
        TEMPLATE_DIR / filename,
        MAPPING_DIR / filename,
        BASE_DIR / filename,
        API_DIR / filename,
    ]

    for path in possible_paths:
        if path.exists():
            return path

    for folder in [TEMPLATE_DIR, MAPPING_DIR, BASE_DIR, API_DIR]:
        if folder.exists():
            for path in folder.rglob(filename):
                if path.exists():
                    return path

    return None


def normalize_list_value(value, allowed_values):
    if is_empty(value):
        return None

    raw = str(value).strip()
    raw_low = raw.lower()

    for allowed in allowed_values:
        if raw == allowed:
            return allowed

    for allowed in allowed_values:
        if raw_low == str(allowed).lower():
            return allowed

    if "animation" in raw_low:
        for allowed in allowed_values:
            if "animation" in str(allowed).lower():
                return allowed

    if "restaurant" in raw_low or "restauration" in raw_low:
        for allowed in allowed_values:
            if "restaurant" in str(allowed).lower():
                return allowed

    if "hébergement" in raw_low or "hebergement" in raw_low or "hotel" in raw_low:
        for allowed in allowed_values:
            a = str(allowed).lower()
            if "hébergement" in a or "hebergement" in a:
                return allowed

    if "création" in raw_low or "creation" in raw_low:
        for allowed in allowed_values:
            a = str(allowed).lower()
            if "création" in a or "creation" in a:
                return allowed

    if "extension" in raw_low:
        for allowed in allowed_values:
            if "extension" in str(allowed).lower():
                return allowed

    return raw


def normalize_excel_value(value, mapping_item, source_path=None, category=None):
    if is_empty(value):
        if "default_by_category" in mapping_item:
            defaults = mapping_item["default_by_category"]
            value = defaults.get(category, defaults.get("default"))
        elif "default" in mapping_item:
            value = mapping_item.get("default")
        else:
            return None

    value_type = mapping_item.get("type")
    unit = str(mapping_item.get("unit", "")).upper()

    if value_type == "list":
        return normalize_list_value(value, mapping_item.get("allowed_values", []))

    if value_type in ["number", "integer", "float", "money"]:
        number = parse_number(value)

        if number is None:
            return None

        if unit == "KMAD":
            if source_path and str(source_path).endswith("_mad"):
                number = number / 1000
            elif abs(number) >= 100000:
                number = number / 1000

        return number

    return value


# =========================================================
# CONTEXTE PAR DÉFAUT ET ALIASES
# =========================================================

def add_context_defaults(context):
    context = context if isinstance(context, dict) else {}

    sections = [
        "dossier",
        "entreprise",
        "dirigeant",
        "projet",
        "banque",
        "emplois",
        "investissements",
        "financement_pi",
        "financement_expert",
        "financement_checkbox",
        "hypotheses_financieres",
        "dap",
        "cpc_historique",
        "cpc_previsionnel",
        "bilan_historique",
        "bilan_previsionnel",
        "impacts_historique",
        "impacts_previsionnels",
    ]

    for section in sections:
        if section == "investissements":
            context.setdefault(section, {})
        else:
            context[section] = as_dict(context.get(section))

    # Copier quelques champs top-level vers les sections si l'IA les renvoie à plat.
    top_to_section = {
        "raison_sociale": "entreprise.raison_sociale",
        "forme_juridique": "entreprise.forme_juridique",
        "rc": "entreprise.rc",
        "ice": "entreprise.ice",
        "cnss": "entreprise.cnss",
        "capital_social": "entreprise.capital_social",
        "capital_social_mad": "entreprise.capital_social_mad",
        "adresse_siege": "entreprise.adresse_siege",
        "activite": "entreprise.activite",
        "dirigeant": "dirigeant.nom",
        "nom_dirigeant": "dirigeant.nom",
        "cin": "dirigeant.cin",
        "qualite": "dirigeant.qualite",
        "mobile": "dirigeant.mobile",
        "email": "dirigeant.email",
        "objet": "projet.objet",
        "ville_region": "projet.ville_region",
        "investissement_total": "projet.investissement_total",
        "investissement_total_mad": "projet.investissement_total_mad",
        "planning_realisation": "projet.planning_realisation",
        "date_demarrage_prevue": "projet.date_demarrage_prevue",
        "effectif": "projet.effectif",
        "banque_nom": "banque.nom",
        "banque_partenaire": "banque.nom",
        "fonds_propres": "financement_expert.fonds_propres_mad",
        "credit_bancaire": "financement_expert.credit_bancaire_mad",
        "cmt": "financement_expert.credit_bancaire_mad",
        "prime": "financement_expert.prime_istitmar_mad",
    }

    for source, target in top_to_section.items():
        if source in context and context[source] not in [None, "", [], {}]:
            if get_context_value(context, target, "") in ["", None, 0]:
                deep_set(context, target, context[source])

    # Identité dossier
    today_fr = date.today().strftime("%d/%m/%Y")

    entreprise = context["entreprise"]
    dirigeant = context["dirigeant"]
    projet = context["projet"]
    dossier = context["dossier"]
    banque = context["banque"]
    financement = context["financement_expert"]
    emplois = context["emplois"]

    raison = entreprise.get("raison_sociale") or entreprise.get("denomination") or context.get("raison_sociale", "")
    rc = entreprise.get("rc") or entreprise.get("numero_rc") or context.get("rc", "")

    dossier.setdefault("date_dossier", today_fr)
    dossier.setdefault("date_signature", dossier.get("date_dossier") or today_fr)
    dossier.setdefault("lieu_signature", projet.get("ville_region", "").split(",")[0].strip() or "Agadir")

    if not dossier.get("identifiant"):
        if context.get("identifiant"):
            dossier["identifiant"] = context["identifiant"]
        elif rc:
            dossier["identifiant"] = f"INV-{rc}"
        elif raison:
            dossier["identifiant"] = f"INV-{safe_name(raison)}"
        else:
            dossier["identifiant"] = "INV-DOSSIER"

    # Entreprise
    entreprise.setdefault("raison_sociale", raison)
    entreprise.setdefault("denomination", entreprise.get("raison_sociale", ""))
    entreprise.setdefault("numero_rc", entreprise.get("rc", ""))
    entreprise.setdefault("adresse", entreprise.get("adresse_siege", ""))
    entreprise.setdefault("secteur_activite", entreprise.get("activite", ""))
    entreprise.setdefault("activite_selection", entreprise.get("activite", "Animation touristique"))
    entreprise.setdefault("activite", entreprise.get("activite_selection", "Animation touristique"))
    entreprise.setdefault("capital_social_mad", to_number(entreprise.get("capital_social_mad") or entreprise.get("capital_social")))

    # Dirigeant
    dirigeant.setdefault("nom", context.get("nom", ""))
    dirigeant.setdefault("qualite", dirigeant.get("fonction") or "Gérant")
    dirigeant.setdefault("fonction", dirigeant.get("qualite", "Gérant"))
    dirigeant.setdefault("gsm", dirigeant.get("mobile", ""))
    dirigeant.setdefault("fixe", dirigeant.get("telephone_fixe", ""))

    # Projet
    projet.setdefault("objectif", projet.get("objet", "Création"))
    projet.setdefault("objet", projet.get("description", ""))
    projet.setdefault("description", projet.get("objet", ""))
    projet.setdefault("investissement_total_mad", projet.get("investissement_total", 0))
    projet.setdefault("adresse_site", projet.get("adresse_installations", ""))
    projet.setdefault("adresse_installations", projet.get("adresse_site", ""))
    projet.setdefault("secteur", projet.get("branche_activite", "Animation touristique"))
    projet.setdefault("branche_activite", projet.get("secteur", "Animation touristique"))
    projet.setdefault("secteur_activite_projet", projet.get("secteur", "Animation touristique"))
    projet.setdefault("filieres", "Animation touristique")
    projet.setdefault("ecosystemes", "Animation touristique et loisirs")
    projet.setdefault("activites_envisagees", projet.get("objet", "Activités d’animation touristique et de loisirs"))
    projet.setdefault("responsable_nom", dirigeant.get("nom", ""))
    projet.setdefault("responsable_projet", projet.get("responsable_nom", ""))
    projet.setdefault("responsable_mobile", dirigeant.get("mobile", ""))
    projet.setdefault("effectif", emplois.get("emplois_directs", emplois.get("directs", projet.get("effectif", ""))))
    projet.setdefault("emplois_indirects", emplois.get("emplois_indirects", ""))
    projet.setdefault("ca_prevu_annee_1", context.get("ca_prevu_2026", projet.get("ca_prevu_annee_1", 0)))
    projet.setdefault("croissance_ca", projet.get("croissance_ca", 0.15))
    projet.setdefault("facteurs_differenciation", context.get("facteurs_differenciation", "Le projet se distingue par une offre d’animation touristique structurée, une expérience client de qualité, une implantation régionale attractive et une organisation opérationnelle adaptée aux standards du secteur."))
    projet.setdefault("attractivite_touristique", context.get("attractivite_touristique", "Le projet contribue au renforcement de l’attractivité touristique de la destination en diversifiant l’offre de loisirs, en améliorant l’expérience des visiteurs et en favorisant la création d’emplois locaux."))

    # Banque
    banque_nom = banque.get("nom") or banque.get("banque_partenaire") or context.get("banque", "")
    banque.setdefault("nom", banque_nom)
    banque.setdefault("banque_partenaire", banque.get("nom", ""))
    banque.setdefault("forme_juridique", banque.get("forme_juridique_banque") or "Société Anonyme")
    banque.setdefault("capital", banque.get("capital_social") or banque.get("capital_social_banque") or ("2 054 500 000" if "populaire" in banque.get("nom", "").lower() else ""))
    banque.setdefault("siege", banque.get("siege_social") or ("101, Boulevard Zerktouni, Casablanca" if "populaire" in banque.get("nom", "").lower() else ""))

    # Financement : harmonisation MAD/KMAD
    fonds_kmad = first_non_empty(context, [
        "financement_expert.fonds_propres_kmad",
        "financement.fonds_propres_kmad",
        "projet.mode_financement.fonds_propres_kmad",
    ], None)
    fonds_mad = first_non_empty(context, [
        "financement_expert.fonds_propres_mad",
        "financement_expert.fonds_propres",
        "financement.fonds_propres_mad",
        "projet.mode_financement.fonds_propres",
        "projet.mode_financement.autofinancement",
        "fonds_propres",
    ], None)

    credit_kmad = first_non_empty(context, [
        "financement_expert.credit_bancaire_kmad",
        "financement_expert.cmt_kmad",
        "financement.cmt_kmad",
        "financement.credit_bancaire_kmad",
        "projet.mode_financement.credit_bancaire_kmad",
    ], None)
    credit_mad = first_non_empty(context, [
        "financement_expert.credit_bancaire_mad",
        "financement_expert.credit_bancaire",
        "financement_expert.cmt",
        "financement.cmt_mad",
        "financement.credit_bancaire_mad",
        "projet.mode_financement.credit_bancaire",
        "projet.mode_financement.cmt",
        "credit_bancaire",
        "cmt",
    ], None)

    prime_kmad = first_non_empty(context, [
        "financement_expert.prime_istitmar_kmad",
        "financement_expert.prime_kmad",
        "prime_deblocage.total_kmad",
    ], None)
    prime_mad = first_non_empty(context, [
        "financement_expert.prime_istitmar_mad",
        "financement_expert.prime_istitmar",
        "financement_expert.prime",
        "financement.prime_mad",
        "projet.mode_financement.prime",
        "prime",
        "montant_prime",
    ], None)

    if fonds_kmad not in [None, ""]:
        financement["fonds_propres_kmad"] = to_number(fonds_kmad)
        financement["fonds_propres_mad"] = to_number(fonds_kmad) * 1000
    elif fonds_mad not in [None, ""]:
        financement["fonds_propres_mad"] = number_as_mad(fonds_mad)
        financement["fonds_propres_kmad"] = number_as_kmad(fonds_mad)

    if credit_kmad not in [None, ""]:
        financement["credit_bancaire_kmad"] = to_number(credit_kmad)
        financement["credit_bancaire_mad"] = to_number(credit_kmad) * 1000
    elif credit_mad not in [None, ""]:
        financement["credit_bancaire_mad"] = number_as_mad(credit_mad)
        financement["credit_bancaire_kmad"] = number_as_kmad(credit_mad)

    if prime_kmad not in [None, ""]:
        financement["prime_istitmar_kmad"] = to_number(prime_kmad)
        financement["prime_istitmar_mad"] = to_number(prime_kmad) * 1000
    elif prime_mad not in [None, ""]:
        financement["prime_istitmar_mad"] = number_as_mad(prime_mad)
        financement["prime_istitmar_kmad"] = number_as_kmad(prime_mad)

    for key in ["financement_participatif", "credit_fournisseur", "leasing"]:
        v_kmad = first_non_empty(context, [f"financement_expert.{key}_kmad", f"financement.{key}_kmad"], None)
        v_mad = first_non_empty(context, [f"financement_expert.{key}_mad", f"financement_expert.{key}", f"financement.{key}_mad"], None)
        if v_kmad not in [None, ""]:
            financement[f"{key}_kmad"] = to_number(v_kmad)
            financement[f"{key}_mad"] = to_number(v_kmad) * 1000
        elif v_mad not in [None, ""]:
            financement[f"{key}_mad"] = number_as_mad(v_mad)
            financement[f"{key}_kmad"] = number_as_kmad(v_mad)
        else:
            financement.setdefault(f"{key}_kmad", 0)
            financement.setdefault(f"{key}_mad", 0)

    # Checkboxes DAP
    context["financement_checkbox"]["autofinancement"] = "☑" if to_number(financement.get("fonds_propres_kmad", 0)) > 0 else "☐"
    context["financement_checkbox"]["cmt"] = "☑" if to_number(financement.get("credit_bancaire_kmad", 0)) > 0 else "☐"
    context["financement_checkbox"]["financement_participatif"] = "☑" if to_number(financement.get("financement_participatif_kmad", 0)) > 0 else "☐"
    context["financement_checkbox"]["credit_fournisseur"] = "☑" if to_number(financement.get("credit_fournisseur_kmad", 0)) > 0 else "☐"
    context["financement_checkbox"]["leasing"] = "☑" if to_number(financement.get("leasing_kmad", 0)) > 0 else "☐"

    # Valeurs narratives DAP
    context.setdefault("synthese_etude_marche", "Le marché ciblé présente un potentiel favorable grâce à la dynamique touristique régionale, à la demande croissante pour des expériences de loisirs structurées et à la complémentarité du projet avec l’offre touristique existante.")
    context.setdefault("facteurs_differenciation", projet.get("facteurs_differenciation", ""))
    context.setdefault("attractivite_touristique", projet.get("attractivite_touristique", ""))

    return context


# Compatibilité avec les anciens appels du fichier.
def add_docx_defaults(context):
    return add_context_defaults(context)


# =========================================================
# BP EXCEL AVANCÉ
# =========================================================

def build_candidate_paths(mapping_item):
    paths = []

    for path in mapping_item.get("field_candidates", []):
        if path and path not in paths:
            paths.append(path)

    field = mapping_item.get("field")
    if field and field not in paths:
        paths.append(field)

    extra = []

    for path in paths:
        if path.endswith("_kmad"):
            extra.append(path[:-5] + "_mad")
        if path.endswith("_mad"):
            extra.append(path[:-4] + "_kmad")
        if ".financement_expert." in path:
            extra.append(path.replace(".financement_expert.", ".financement."))
        if path.startswith("financement_expert."):
            extra.append(path.replace("financement_expert.", "financement."))
        if path.startswith("banque."):
            extra.append(path.replace("banque.", "projet.banque."))
        if path.startswith("entreprise."):
            extra.append(path.split(".", 1)[1])
        if path.startswith("projet."):
            extra.append(path.split(".", 1)[1])

    for path in extra:
        if path not in paths:
            paths.append(path)

    return paths


def get_value_from_mapping(data, mapping_item):
    for path in build_candidate_paths(mapping_item):
        value = get_context_value(data, path, None)

        if value not in [None, "", [], {}]:
            return value, path

    if "default" in mapping_item:
        return mapping_item.get("default"), "__default__"

    return None, None


def cell_to_row_col(cell_ref):
    match = re.match(r"^([A-Z]+)([0-9]+)$", str(cell_ref).upper())

    if not match:
        return None, None

    col = column_index_from_string(match.group(1))
    row = int(match.group(2))

    return row, col


def build_blocked_ranges(mapping):
    blocked = {}

    for rule in mapping.get("never_write", []):
        sheet = rule.get("sheet")

        if not sheet:
            continue

        blocked.setdefault(sheet, [])

        for rng in rule.get("ranges", []):
            try:
                min_col, min_row, max_col, max_row = range_boundaries(rng)
                blocked[sheet].append((min_row, min_col, max_row, max_col, rng))
            except Exception:
                pass

    return blocked


def is_blocked_cell(sheet_name, cell_ref, blocked_ranges):
    row, col = cell_to_row_col(cell_ref)

    if row is None:
        return True

    for min_row, min_col, max_row, max_col, _ in blocked_ranges.get(sheet_name, []):
        if min_row <= row <= max_row and min_col <= col <= max_col:
            return True

    return False


def write_excel_cell(ws, cell_ref, value, blocked_ranges=None, overwrite_formula=False):
    blocked_ranges = blocked_ranges or {}

    if is_blocked_cell(ws.title, cell_ref, blocked_ranges):
        return False

    try:
        cell = ws[cell_ref]
    except Exception:
        return False

    if not overwrite_formula:
        if getattr(cell, "data_type", None) == "f" or (isinstance(cell.value, str) and cell.value.startswith("=")):
            return False

    if value is None:
        return False

    try:
        cell.value = value
        return True
    except Exception:
        return False


def write_mapping_section(wb, data, mapping, section_name, blocked_ranges):
    count = 0

    for item in mapping.get(section_name, []):
        sheet_name = item.get("sheet")
        cell_ref = item.get("cell")

        if not sheet_name or not cell_ref:
            continue

        if sheet_name not in wb.sheetnames:
            continue

        value, source_path = get_value_from_mapping(data, item)
        value = normalize_excel_value(value, item, source_path=source_path)

        if value is None:
            continue

        if write_excel_cell(wb[sheet_name], cell_ref, value, blocked_ranges):
            count += 1

    return count


def normalize_category_name(value):
    text = normalize_key_text(value)

    if "terrain" in text:
        return "terrain"
    if "construction" in text:
        return "constructions"
    if "amenagement" in text or "agencement" in text:
        return "amenagement_agencement"
    if "materiel" in text or "equipement" in text:
        return "materiel_equipement"
    if "frais" in text and ("prelim" in text or "prealable" in text):
        return "frais_preliminaires"
    if "divers" in text or "imprevu" in text:
        return "divers_imprevus"

    return text


def normalize_investissements(data):
    investissements = data.get("investissements", {})

    result = {
        "terrain": [],
        "constructions": [],
        "amenagement_agencement": [],
        "materiel_equipement": [],
        "frais_preliminaires": [],
        "divers_imprevus": [],
        "frais_approche": [],
    }

    if isinstance(investissements, dict):
        for key, value in investissements.items():
            norm = normalize_category_name(key)
            if norm in result:
                if isinstance(value, list):
                    result[norm] = value
                elif isinstance(value, dict):
                    result[norm] = [value]
            elif isinstance(value, list):
                result.setdefault(norm, value)
        return result

    if isinstance(investissements, list):
        for item in investissements:
            if not isinstance(item, dict):
                continue

            cat = normalize_category_name(
                item.get("categorie")
                or item.get("category")
                or item.get("type")
                or item.get("rubrique")
                or item.get("nature")
                or ""
            )

            if cat not in result:
                cat = "materiel_equipement"

            result[cat].append(item)

    return result


def get_bp_source_array(data, source_array, category):
    arr = get_context_value(data, source_array, None)

    if isinstance(arr, list):
        return arr

    if isinstance(arr, dict):
        return list(arr.values())

    investissements = normalize_investissements(data)
    return investissements.get(category, [])


def get_default_tva(col_map, category):
    if "default_by_category" in col_map:
        defaults = col_map["default_by_category"]
        value = defaults.get(category, defaults.get("default"))
        if value is not None:
            return to_number(value)

    if "default" in col_map:
        return to_number(col_map.get("default"))

    return 0.2 if category != "terrain" else 0


def get_investment_cell_value(row_data, col_map, category):
    field = col_map.get("field")
    value = deep_get(row_data, field, None) if field else None

    # Le BP attend le montant HT dans la colonne K.
    # Si l'IA donne seulement un montant global/TTC, on convertit en HT afin que le total TTC reste cohérent.
    if field == "montant_ht_devise" and value in [None, "", [], {}]:
        candidate_fields = [
            "montant_ht_devise",
            "montant_ht",
            "montant_ht_mad",
            "montant_ttc",
            "montant_ttc_mad",
            "montant",
            "montant_mad",
            "cout_mad",
            "total_mad",
            "valeur",
            "prix",
        ]

        source_field = None
        for cand in candidate_fields:
            v = deep_get(row_data, cand, None)
            if v not in [None, "", [], {}]:
                value = v
                source_field = cand
                break

        if source_field and "ht" not in source_field:
            taux = deep_get(row_data, "taux_tva", None)
            if taux in [None, ""]:
                taux = get_default_tva(col_map, category)
            taux = to_number(taux)
            if taux > 1:
                taux = taux / 100
            if taux > 0:
                value = to_number(value) / (1 + taux)

    if value in [None, "", [], {}]:
        if "default_by_category" in col_map:
            defaults = col_map["default_by_category"]
            value = defaults.get(category, defaults.get("default"))
        elif "default" in col_map:
            value = col_map["default"]

    return value


def write_bp_table_mappings(wb, data, mapping, blocked_ranges):
    count = 0

    for table in mapping.get("table_mappings", []):
        sheet_name = table.get("sheet")
        category = table.get("category")
        start_row = table.get("start_row")
        end_row = table.get("end_row")

        if sheet_name not in wb.sheetnames:
            continue

        if start_row is None or end_row is None:
            continue

        ws = wb[sheet_name]
        rows = get_bp_source_array(data, table.get("source_array"), category)

        if not isinstance(rows, list):
            continue

        max_rows = max(0, int(end_row) - int(start_row) + 1)

        for index, row_data in enumerate(rows[:max_rows]):
            excel_row = int(start_row) + index

            for col_map in table.get("columns", []):
                col_letter = col_map.get("column")
                field = col_map.get("field")

                if not col_letter or not field:
                    continue

                value = get_investment_cell_value(row_data, col_map, category)

                if value in [None, "", [], {}]:
                    continue

                value = normalize_excel_value(value, col_map, source_path=field, category=category)
                cell_ref = f"{col_letter}{excel_row}"

                if write_excel_cell(ws, cell_ref, value, blocked_ranges):
                    count += 1

    return count


def force_excel_recalculation(wb):
    try:
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
        wb.calculation.calcMode = "auto"
    except Exception:
        pass


def apply_bp_post_corrections(wb, context):
    """
    Corrections BP après mapping :
    - Prime ISTITMAR sans #NAME?
    - Financement en KMAD cohérent
    - Report à nouveau chaîné
    - Trésorerie actif/passif comme variable d'équilibrage du Bilan
    """
    if "Mode de financement" in wb.sheetnames:
        ws = wb["Mode de financement"]

        fonds_kmad = to_number(get_context_value(context, "financement_expert.fonds_propres_kmad", 0))
        credit_kmad = to_number(get_context_value(context, "financement_expert.credit_bancaire_kmad", 0))
        fp_kmad = to_number(get_context_value(context, "financement_expert.financement_participatif_kmad", 0))
        cf_kmad = to_number(get_context_value(context, "financement_expert.credit_fournisseur_kmad", 0))
        leasing_kmad = to_number(get_context_value(context, "financement_expert.leasing_kmad", 0))
        prime_kmad = to_number(get_context_value(context, "financement_expert.prime_istitmar_kmad", 0))

        if not prime_kmad:
            prime_mad = to_number(get_context_value(context, "financement_expert.prime_istitmar_mad", 0))
            prime_kmad = prime_mad / 1000 if prime_mad else 0

        # Cellules principales du mode de financement
        if fonds_kmad:
            write_excel_cell(ws, "G16", fonds_kmad, overwrite_formula=True)
        if credit_kmad:
            write_excel_cell(ws, "H16", credit_kmad, overwrite_formula=True)
        if fp_kmad:
            write_excel_cell(ws, "I16", fp_kmad, overwrite_formula=True)
        if cf_kmad:
            write_excel_cell(ws, "J16", cf_kmad, overwrite_formula=True)
        if leasing_kmad:
            write_excel_cell(ws, "K16", leasing_kmad, overwrite_formula=True)
        if prime_kmad:
            # L16 est souvent la prime ; l'écriture supprime le #NAME?.
            write_excel_cell(ws, "L16", prime_kmad, overwrite_formula=True)

        for target_col, source_col in zip(["N", "O", "P", "Q", "R", "S"], ["G", "H", "I", "J", "K", "L"]):
            write_excel_cell(ws, f"{target_col}16", f"=IFERROR({source_col}16/SUM($G$16:$L$16),0)", overwrite_formula=True)

    if "Bilan" in wb.sheetnames:
        ws = wb["Bilan"]

        # Chaînage du report à nouveau / résultats non affectés
        write_excel_cell(ws, "H56", 0, overwrite_formula=True)
        write_excel_cell(ws, "I56", "=H56+H57", overwrite_formula=True)
        write_excel_cell(ws, "J56", "=I56+I57", overwrite_formula=True)
        write_excel_cell(ws, "K56", "=J56+J57", overwrite_formula=True)
        write_excel_cell(ws, "L56", "=K56+K57", overwrite_formula=True)
        write_excel_cell(ws, "M56", "=L56+L57", overwrite_formula=True)

        # Variable de bouclage : trésorerie actif/passif
        for col in ["H", "I", "J", "K", "L", "M"]:
            write_excel_cell(ws, f"{col}45", f"=MAX(0,({col}68+{col}78)-({col}27+{col}43))", overwrite_formula=True)
            write_excel_cell(ws, f"{col}80", f"=MAX(0,({col}27+{col}43)-({col}68+{col}78))", overwrite_formula=True)
            write_excel_cell(ws, f"{col}85", f"={col}47-{col}82", overwrite_formula=True)

        write_excel_cell(ws, "A3", '=IF(AND(ABS(H85)<50,ABS(I85)<50,ABS(J85)<50,ABS(K85)<50,ABS(L85)<50,ABS(M85)<50),"Equilibré","Déséquilibré")', overwrite_formula=True)


def render_bp_excel(output_path, context):
    template_path = find_file("BP_template.xlsx")
    mapping_path = find_file("mapping_bp_istitmar.json")

    if template_path is None:
        raise FileNotFoundError("BP_template.xlsx introuvable dans /templates")

    if mapping_path is None:
        raise FileNotFoundError("mapping_bp_istitmar.json introuvable dans /mappings")

    with open(mapping_path, "r", encoding="utf-8") as f:
        mapping = json.load(f)

    wb = load_workbook(template_path)
    blocked_ranges = build_blocked_ranges(mapping)

    written = {
        "scalar_mappings": write_mapping_section(wb, context, mapping, "scalar_mappings", blocked_ranges),
        "financement_pi_mappings": write_mapping_section(wb, context, mapping, "financement_pi_mappings", blocked_ranges),
        "table_mappings": write_bp_table_mappings(wb, context, mapping, blocked_ranges),
        "cpc_mappings": write_mapping_section(wb, context, mapping, "cpc_mappings", blocked_ranges),
        "bilan_mappings": write_mapping_section(wb, context, mapping, "bilan_mappings", blocked_ranges),
        "impacts_mappings": write_mapping_section(wb, context, mapping, "impacts_mappings", blocked_ranges),
    }

    apply_bp_post_corrections(wb, context)
    force_excel_recalculation(wb)
    wb.save(output_path)

    return written


# =========================================================
# WORD / OUTILS ROBUSTES
# =========================================================

def normalize_label(text):
    text = str(text or "").lower().replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\wÀ-ÿ%()./ -]", "", text)
    return text.strip()


def safe_table_rows(table):
    try:
        return list(table.rows)
    except Exception:
        return []


def safe_row_cells(row, table):
    cells = []

    try:
        for child in row._tr.iterchildren():
            if child.tag == qn("w:tc"):
                cells.append(_Cell(child, row))
            elif child.tag == qn("w:sdt"):
                for sdt_child in child.iter():
                    if sdt_child.tag == qn("w:tc"):
                        cells.append(_Cell(sdt_child, row))
    except Exception:
        pass

    if cells:
        return cells

    try:
        return list(row.cells)
    except Exception:
        try:
            return [_Cell(tc, row) for tc in row._tr.tc_lst]
        except Exception:
            return []


def get_table_by_index(doc, index):
    try:
        return doc.tables[int(index)]
    except Exception:
        return None


def iter_container_paragraphs(container):
    try:
        for paragraph in container.paragraphs:
            yield paragraph
    except Exception:
        pass

    try:
        for table in container.tables:
            rows = safe_table_rows(table)
            for row in rows:
                cells = safe_row_cells(row, table)
                for cell in cells:
                    try:
                        for paragraph in cell.paragraphs:
                            yield paragraph
                    except Exception:
                        continue
    except Exception:
        pass


def iter_all_paragraphs(doc):
    yield from iter_container_paragraphs(doc)

    try:
        for section in doc.sections:
            yield from iter_container_paragraphs(section.header)
            yield from iter_container_paragraphs(section.footer)
    except Exception:
        pass


def set_cell_value(cell, value):
    try:
        cell.text = "" if value is None else str(value)
    except Exception:
        pass


def set_paragraph_text(paragraph, value):
    try:
        paragraph.text = "" if value is None else str(value)
    except Exception:
        pass


def replace_text_in_xml_container(container, search, replace):
    if not search:
        return

    replace = "" if replace is None else str(replace)

    try:
        root = container.element
    except Exception:
        try:
            root = container._element
        except Exception:
            return

    try:
        for node in root.iter(qn("w:t")):
            if node.text and search in node.text:
                node.text = node.text.replace(search, replace)
    except Exception:
        pass


def replace_text_everywhere(doc, search, replace):
    if not search:
        return

    replace = "" if replace is None else str(replace)

    for paragraph in iter_all_paragraphs(doc):
        try:
            if search in paragraph.text:
                paragraph.text = paragraph.text.replace(search, replace)
        except Exception:
            continue

    replace_text_in_xml_container(doc, search, replace)

    try:
        for section in doc.sections:
            replace_text_in_xml_container(section.header, search, replace)
            replace_text_in_xml_container(section.footer, search, replace)
    except Exception:
        pass


def fill_cell_next_to_label(doc, label, value, position="right", occurrence=1):
    wanted = normalize_label(label)
    seen = 0

    for table in doc.tables:
        rows = safe_table_rows(table)

        for r_idx, row in enumerate(rows):
            cells = safe_row_cells(row, table)

            for c_idx, cell in enumerate(cells):
                try:
                    cell_text = normalize_label(cell.text)
                except Exception:
                    continue

                if wanted and wanted in cell_text:
                    seen += 1

                    if seen < occurrence:
                        continue

                    if position == "below" and r_idx + 1 < len(rows):
                        below_cells = safe_row_cells(rows[r_idx + 1], table)

                        if c_idx < len(below_cells):
                            set_cell_value(below_cells[c_idx], value)
                            return True

                    offset = 1

                    if isinstance(position, str) and position.startswith("right"):
                        try:
                            offset = int(position.replace("right", "") or "1")
                        except Exception:
                            offset = 1

                    target_index = c_idx + offset

                    if target_index < len(cells):
                        set_cell_value(cells[target_index], value)
                        return True

    return False


# =========================================================
# DAP WORD AVANCÉ
# =========================================================

def get_expert_default(mapping, field):
    defaults = mapping.get("expert_generation_defaults", {})

    if field in defaults:
        return defaults[field]

    default_texts = {
        "synthese_etude_marche": "Le projet répond à une demande croissante pour des expériences touristiques différenciées et structurées. Il s’inscrit dans une dynamique régionale favorable portée par la reprise du tourisme, la diversification de l’offre de loisirs et l’intérêt des visiteurs pour des activités expérientielles.",
        "facteurs_differenciation": "Le projet se différencie par la qualité de l’expérience proposée, la structuration de l’offre, la sécurité des prestations, l’ancrage régional et la capacité à créer des partenariats avec les acteurs touristiques locaux.",
        "attractivite_touristique": "Le projet contribue à l’attractivité de la destination Maroc en diversifiant les activités de loisirs, en améliorant l’expérience client et en renforçant l’offre d’animation touristique disponible au niveau régional.",
    }

    return default_texts.get(field, "")


def generate_default_rows(source_array, context):
    key = str(source_array).split(".")[-1]

    projet = context.get("projet", {})
    emplois = context.get("emplois", {})

    objet = projet.get("objet") or projet.get("description") or "activité d’animation touristique"
    ville = projet.get("ville_region") or projet.get("region") or "la zone d’implantation"
    ca = to_number(projet.get("ca_prevu_annee_1", 0))
    croissance = to_number(projet.get("croissance_ca", 0.15)) or 0.15
    effectif = int(to_number(projet.get("effectif") or emplois.get("directs") or emplois.get("emplois_directs") or 5))

    if key == "gamme_services":
        return [
            {
                "domaine_activite": "Animation touristique",
                "description": objet,
                "marche_adressable": f"Touristes nationaux, touristes étrangers et clientèle locale à {ville}",
                "pourcentage_ca": "100%",
                "image_slot": "Image illustrative à ajouter",
            }
        ]

    if key == "marche_cibles":
        return [
            {
                "domaine_activite": "Animation touristique",
                "marche_cible": "Touristes nationaux, touristes internationaux, familles, groupes et clientèle locale",
                "taille_marche_mmad": "",
                "tcam": "",
                "source_taille_marche": "Estimation interne",
                "source_tcam": "Estimation interne",
            }
        ]

    if key == "concurrents":
        return [
            {
                "nom": "Opérateurs touristiques locaux",
                "implantation": ville,
                "categorie": "Concurrence directe et indirecte",
                "part_marche": "",
                "principaux_services": "Activités touristiques, loisirs, animation et expériences clients",
            }
        ]

    if key == "fournisseurs":
        return [
            {
                "nom": "Fournisseurs locaux et nationaux",
                "categorie": "Équipements, consommables et maintenance",
                "produits_services": "Matériel d’exploitation, maintenance, fournitures et prestations de support",
                "part_achats": "",
                "delai_reglement_jours": "30",
                "modalites_achat": "Achat sur devis, bon de commande et règlement selon conditions négociées",
            }
        ]

    if key == "clients_creneaux":
        return [
            {
                "creneau_service": "Clientèle touristique",
                "principaux_clients": "Touristes nationaux, touristes étrangers, familles, groupes et agences partenaires",
                "categorie": "B2C/B2B",
                "part_ca": "100%",
                "delai_recouvrement_jours": "0",
                "modalites_vente": "Paiement comptant, réservation en ligne et partenariats commerciaux",
            }
        ]

    if key == "politique_prix":
        return [
            {
                "famille_service": "Animation touristique",
                "strategie_prix": "Prix aligné sur le marché avec différenciation par la qualité de l’expérience",
                "commentaire": "Tarification modulable selon la saison, les groupes, les offres packagées et les partenariats",
            }
        ]

    if key == "capacites_animation":
        return [
            {
                "domaine_activite": "Animation touristique",
                "capacite_journaliere": "À préciser",
                "duree_cycle_service": "À préciser",
                "prix_moyen_service": "À préciser",
                "taux_occupation": "50%",
            }
        ]

    if key == "emplois_directs_profils":
        return [
            {
                "profil_domaine": "Exploitation et animation",
                "emplois_permanents": effectif,
                "emplois_feminins": "",
                "contrat_cdi": effectif,
                "contrat_anapec": "",
            }
        ]

    if key == "lignes":
        return [
            {
                "libelle": "Animation touristique",
                "taux_croissance": format_display_value(croissance, "percent"),
                "2026_kmad": round(ca / 1000) if ca else "",
                "2027_kmad": round(ca * (1 + croissance) / 1000) if ca else "",
                "2028_kmad": round(ca * (1 + croissance) ** 2 / 1000) if ca else "",
                "2029_kmad": round(ca * (1 + croissance) ** 3 / 1000) if ca else "",
                "2030_kmad": round(ca * (1 + croissance) ** 4 / 1000) if ca else "",
                "2031_kmad": round(ca * (1 + croissance) ** 5 / 1000) if ca else "",
            }
        ]

    return []


def get_dap_source_array(context, source_array, mapping_item=None):
    data = get_context_value(context, source_array, None)

    if isinstance(data, list):
        return data

    if isinstance(data, dict):
        return [data]

    if mapping_item and mapping_item.get("generate_defaults_if_empty"):
        return generate_default_rows(source_array, context)

    generated = generate_default_rows(source_array, context)

    if generated:
        return generated

    return []


def apply_dap_cover_replacements(doc, context):
    dossier = context.get("dossier", {})
    entreprise = context.get("entreprise", {})

    identifiant = dossier.get("identifiant", "INV-DOSSIER")
    raison = entreprise.get("raison_sociale", "SOCIETE")
    date_doc = dossier.get("date_dossier") or dossier.get("date_signature") or date.today().strftime("%d/%m/%Y")

    replacements = {
        "INV-XXX | MaSociété": f"{identifiant} | {raison}",
        "INV-XXX": identifiant,
        "MaSociété": raison,
        "Click or tap to enter a date.": date_doc,
        "Click or tap to enter a date": date_doc,
    }

    for old, new in replacements.items():
        replace_text_everywhere(doc, old, new)

    for paragraph in iter_all_paragraphs(doc):
        try:
            txt = paragraph.text
            if "INV-" in txt and "MaSociété" in txt:
                paragraph.text = f"{identifiant} | {raison}"
            elif "Click or tap to enter" in txt:
                paragraph.text = date_doc
        except Exception:
            continue


def apply_scalar_text_replacements(doc, context, mapping):
    for item in mapping.get("scalar_text_replacements", []):
        placeholder = item.get("placeholder") or item.get("search")
        field = item.get("field")
        fmt = item.get("format") or item.get("type")

        if not placeholder:
            continue

        value = get_context_value(context, field, item.get("default", "")) if field else item.get("replacement", "")

        if value in [None, ""]:
            value = item.get("default", "")

        value = format_display_value(value, fmt)
        replace_text_everywhere(doc, placeholder, value)


def apply_legacy_text_replacements(doc, context, mapping):
    for item in mapping.get("text_replacements", []) + mapping.get("literal_replacements", []):
        search = item.get("search")
        field = item.get("field")
        replacement = item.get("replacement")

        if not search:
            continue

        if field:
            replacement = get_context_value(context, field, "")

        replace_text_everywhere(doc, search, replacement or "")


def apply_legacy_label_mappings(doc, context, mapping):
    label_mappings = (
        mapping.get("cell_mappings", [])
        + mapping.get("label_mappings", [])
        + mapping.get("field_mappings", [])
    )

    for item in label_mappings:
        label = item.get("label") or item.get("search")
        field = item.get("field")
        position = item.get("position", "right")
        occurrence = int(item.get("occurrence", 1))
        fmt = item.get("format") or item.get("type")

        if not label or not field:
            continue

        value = get_context_value(context, field, "")

        if value in ["", None]:
            continue

        fill_cell_next_to_label(doc, label, format_display_value(value, fmt), position, occurrence)


def apply_table_cell_mappings(doc, context, mapping):
    for item in mapping.get("table_cell_mappings", []):
        table = get_table_by_index(doc, item.get("table_index"))

        if table is None:
            continue

        rows = safe_table_rows(table)
        row_idx = int(item.get("row", 0))

        if row_idx < 0 or row_idx >= len(rows):
            continue

        cells = safe_row_cells(rows[row_idx], table)

        if not cells:
            continue

        target = item.get("target", "right")
        fmt = item.get("format") or item.get("type")

        if target == "split_cells":
            fields = item.get("fields", [])

            for i, field in enumerate(fields):
                cell_idx = i + 1

                if cell_idx < len(cells):
                    value = get_context_value(context, field, "")
                    set_cell_value(cells[cell_idx], format_display_value(value, fmt))

            continue

        if target == "third_cell_append":
            field = item.get("field")
            value = get_context_value(context, field, "")

            if len(cells) >= 3 and value not in ["", None]:
                old = cells[2].text.strip()
                new_text = (old + " " + format_display_value(value, fmt)).strip()
                set_cell_value(cells[2], new_text)

            continue

        if target == "authorization_row":
            array_path = item.get("array")
            array_index = int(item.get("array_index", 0))
            rows_data = get_dap_source_array(context, array_path)

            if 0 <= array_index < len(rows_data):
                row_data = rows_data[array_index]

                if isinstance(row_data, dict):
                    value = " | ".join(str(v) for v in row_data.values() if v not in ["", None])
                else:
                    value = str(row_data)

                if len(cells) > 1:
                    set_cell_value(cells[1], value)

            continue

        field = item.get("field")
        value = get_context_value(context, field, "")

        if item.get("generate_if_empty") and value in ["", None]:
            value = get_expert_default(mapping, field)

        value = format_display_value(value, fmt)

        # Important : ne jamais écrire dans la cellule libellé si cellule de droite absente.
        if len(cells) < 2:
            continue

        set_cell_value(cells[1], value)


def fill_image_cell(cell, image_path, fallback_text, width_cm=4.5, height_cm=3.0):
    if image_path:
        p = Path(str(image_path))

        if p.exists():
            try:
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(str(p), width=Cm(float(width_cm)), height=Cm(float(height_cm)))
                return
            except Exception:
                pass

    set_cell_value(cell, fallback_text or "Image illustrative à ajouter")


def apply_repeat_table_mappings(doc, context, mapping):
    image_settings_by_table = {}

    for img in mapping.get("image_mappings", []):
        if "table_index" in img:
            image_settings_by_table[int(img["table_index"])] = img

    for item in mapping.get("repeat_table_mappings", []):
        table = get_table_by_index(doc, item.get("table_index"))

        if table is None:
            continue

        rows = safe_table_rows(table)
        source_array = item.get("source_array", "")
        data_rows = get_dap_source_array(context, source_array, item)

        if not data_rows:
            continue

        start_row = int(item.get("start_row", 0))
        row_step = int(item.get("row_step", 1))
        max_rows = int(item.get("max_rows", len(data_rows)))
        columns = item.get("columns", [])

        for i, row_data in enumerate(data_rows[:max_rows]):
            row_idx = start_row + i * row_step

            if row_idx >= len(rows):
                break

            cells = safe_row_cells(rows[row_idx], table)

            for col in columns:
                col_idx = col.get("col", col.get("index", col.get("column_index")))
                field = col.get("field")
                fmt = col.get("format") or col.get("type")

                try:
                    col_idx = int(col_idx)
                except Exception:
                    continue

                if col_idx < 0 or col_idx >= len(cells):
                    continue

                if col.get("type") == "image":
                    img_conf = image_settings_by_table.get(int(item.get("table_index", -1)), {})
                    image_field = img_conf.get("image_field", field)
                    image_path = row_data.get(image_field, "")
                    fallback = img_conf.get("fallback_text", "Image illustrative à ajouter")
                    fill_image_cell(
                        cells[col_idx],
                        image_path,
                        fallback,
                        img_conf.get("insert_width_cm", 4.5),
                        img_conf.get("insert_height_cm", 3.0),
                    )
                    continue

                value = row_data.get(field, "")
                set_cell_value(cells[col_idx], format_display_value(value, fmt))

            source_row_columns = item.get("source_row_columns", [])

            if source_row_columns and row_step >= 2 and row_idx + 1 < len(rows):
                source_cells = safe_row_cells(rows[row_idx + 1], table)

                for src in source_row_columns:
                    col_idx = src.get("col")
                    field = src.get("field")
                    prefix = src.get("prefix", "")

                    try:
                        col_idx = int(col_idx)
                    except Exception:
                        continue

                    if col_idx < 0 or col_idx >= len(source_cells):
                        continue

                    value = row_data.get(field, "")

                    if value not in ["", None]:
                        set_cell_value(source_cells[col_idx], prefix + str(value))

        total_row = item.get("total_row")

        if isinstance(total_row, dict):
            row_idx = int(total_row.get("row", 0))

            if 0 <= row_idx < len(rows):
                cells = safe_row_cells(rows[row_idx], table)

                for col in total_row.get("columns", []):
                    col_idx = col.get("col")
                    field = col.get("field")
                    fmt = col.get("format") or col.get("type")

                    try:
                        col_idx = int(col_idx)
                    except Exception:
                        continue

                    if col_idx < 0 or col_idx >= len(cells):
                        continue

                    value = get_context_value(context, field, "")

                    if value not in ["", None]:
                        set_cell_value(cells[col_idx], format_display_value(value, fmt))


def apply_single_row_table_mappings(doc, context, mapping):
    for item in mapping.get("single_row_table_mappings", []):
        table = get_table_by_index(doc, item.get("table_index"))

        if table is None:
            continue

        rows = safe_table_rows(table)

        if not rows:
            continue

        row_idx = int(item.get("row", 1))

        if row_idx >= len(rows):
            row_idx = 0

        cells = safe_row_cells(rows[row_idx], table)

        target_col = item.get("target_col")
        field = item.get("field")

        try:
            target_col = int(target_col)
        except Exception:
            target_col = None

        if target_col is not None and 0 <= target_col < len(cells):
            value = get_context_value(context, field, "")
            set_cell_value(cells[target_col], format_display_value(value, item.get("format")))

        target_col_2 = item.get("target_col_2")
        field_2 = item.get("field_2")

        try:
            target_col_2 = int(target_col_2)
        except Exception:
            target_col_2 = None

        if target_col_2 is not None and 0 <= target_col_2 < len(cells):
            value_2 = get_context_value(context, field_2, "")

            if item.get("generate_field_2_if_empty") and value_2 in ["", None]:
                value_2 = get_expert_default(mapping, field_2)

            set_cell_value(cells[target_col_2], value_2)


def apply_paragraph_mappings(doc, context, mapping):
    for item in mapping.get("paragraph_mappings", []):
        search = item.get("search_contains")
        field = item.get("field")

        if not search or not field:
            continue

        value = get_context_value(context, field, "")

        if item.get("generate_if_empty") and value in ["", None]:
            value = get_expert_default(mapping, field)

        if value in ["", None]:
            continue

        search_norm = normalize_label(search)

        for paragraph in iter_all_paragraphs(doc):
            try:
                if search_norm in normalize_label(paragraph.text):
                    set_paragraph_text(paragraph, value)
            except Exception:
                continue


def set_checkbox_label(doc, label, checked):
    mark = "☑" if checked else "☐"
    label_norm = normalize_label(label)

    for pattern in [f"☐ {label}", f"□ {label}", f"☑ {label}", f"☐{label}", f"□{label}", f"☑{label}"]:
        replace_text_everywhere(doc, pattern, f"{mark} {label}")

    for paragraph in iter_all_paragraphs(doc):
        try:
            txt_norm = normalize_label(paragraph.text)
            if label_norm in txt_norm and ("☐" in paragraph.text or "□" in paragraph.text or "☑" in paragraph.text):
                paragraph.text = f"{mark} {label}"
        except Exception:
            continue


def apply_checkbox_mappings(doc, context, mapping):
    for item in mapping.get("checkbox_mappings", []):
        label = item.get("label")
        field = item.get("field")

        if not label or not field:
            continue

        raw = get_context_value(context, field, "☐")
        checked = raw is True or str(raw).strip() in ["☑", "true", "True", "1", "oui", "Oui", "yes", "checked"]
        set_checkbox_label(doc, label, checked)


def apply_dap_financing_checkboxes(doc, context):
    financing_labels = [
        ("Autofinancement", "financement_checkbox.autofinancement"),
        ("CMT", "financement_checkbox.cmt"),
        ("Financement participatif", "financement_checkbox.financement_participatif"),
        ("Crédit fournisseur", "financement_checkbox.credit_fournisseur"),
        ("Leasing", "financement_checkbox.leasing"),
    ]

    for label, field in financing_labels:
        raw = get_context_value(context, field, "☐")
        checked = str(raw).strip() == "☑"
        set_checkbox_label(doc, label, checked)


def apply_dap_default_mappings(doc, context):
    default_mappings = [
        {"label": "Identifiant", "field": "dossier.identifiant"},
        {"label": "Raison sociale", "field": "entreprise.raison_sociale"},
        {"label": "Forme juridique", "field": "entreprise.forme_juridique"},
        {"label": "Date de création", "field": "entreprise.date_creation"},
        {"label": "Secteur d’activité", "field": "entreprise.activite"},
        {"label": "Type et Catégorie", "field": "entreprise.type_categorie"},
        {"label": "Capital social (MAD)", "field": "entreprise.capital_social_mad"},
        {"label": "Actionnaires (%)", "field": "entreprise.actionnaires"},
        {"label": "Registre de commerce", "field": "entreprise.rc"},
        {"label": "Identifiant Commun de l'Entreprise", "field": "entreprise.ice"},
        {"label": "CNSS", "field": "entreprise.cnss"},
        {"label": "Adresse du siège social", "field": "entreprise.adresse_siege"},
        {"label": "Adresse du site d’implantation", "field": "projet.adresse_installations"},
        {"label": "Dirigeant", "field": "dirigeant.nom"},
        {"label": "Mobile (Cellulaire)", "field": "dirigeant.mobile"},
        {"label": "Courrier électronique", "field": "dirigeant.email"},
        {"label": "Filière(s) du projet", "field": "projet.filieres"},
        {"label": "Objet du projet", "field": "projet.objet"},
        {"label": "Adresse des installations", "field": "projet.adresse_installations"},
        {"label": "Ville/Région du Projet", "field": "projet.ville_region"},
        {"label": "Superficie et mode d’occupation du site", "field": "projet.surface_mode_occupation"},
        {"label": "Effectif à embaucher", "field": "projet.effectif"},
        {"label": "Planning de réalisation", "field": "projet.planning_realisation"},
        {"label": "Date prévue de démarrage", "field": "projet.date_demarrage_prevue"},
        {"label": "Responsable de projet", "field": "projet.responsable_projet"},
        {"label": "Mobile (Cellulaire) du resp. de projet", "field": "projet.responsable_mobile"},
        {"label": "Partenaire financier", "field": "banque.nom"},
        {"label": "Investissement total", "field": "projet.investissement_total"},
        {"label": "Secteur d’activité du projet", "field": "projet.secteur_activite_projet"},
        {"label": "Ecosystème(s) du projet", "field": "projet.ecosystemes"},
        {"label": "Activité(s) envisagée(s)", "field": "projet.activites_envisagees"},
        {"label": "Fiche(s) de projet", "field": "projet.fiches_projet"},
    ]

    for item in default_mappings:
        value = get_context_value(context, item["field"], "")

        if value not in ["", None, 0]:
            fill_cell_next_to_label(doc, item["label"], value)


def apply_dap_narrative_fallbacks(doc, context):
    narrative_rules = [
        {
            "search": "Synthétiser l’étude de marché réalisée",
            "value": get_context_value(context, "synthese_etude_marche", ""),
        },
        {
            "search": "Décrire brièvement les facteurs de différenciation",
            "value": get_context_value(context, "facteurs_differenciation", ""),
        },
        {
            "search": "Décrire succinctement l’attractivité touristique estimée",
            "value": get_context_value(context, "attractivite_touristique", ""),
        },
    ]

    for rule in narrative_rules:
        search_norm = normalize_label(rule["search"])
        value = rule["value"]

        if not value:
            continue

        for paragraph in iter_all_paragraphs(doc):
            try:
                if search_norm in normalize_label(paragraph.text):
                    paragraph.text = value
                    break
            except Exception:
                continue


def apply_dap_mapping_file(doc, context):
    mapping_path = find_file("mapping_dap_istitmar.json")

    if mapping_path is None:
        apply_dap_cover_replacements(doc, context)
        apply_dap_default_mappings(doc, context)
        apply_dap_financing_checkboxes(doc, context)
        apply_dap_narrative_fallbacks(doc, context)
        return {"mapping_used": False, "reason": "mapping_dap_istitmar.json introuvable"}

    with open(mapping_path, "r", encoding="utf-8") as f:
        mapping = json.load(f)

    apply_dap_cover_replacements(doc, context)
    apply_scalar_text_replacements(doc, context, mapping)
    apply_legacy_text_replacements(doc, context, mapping)
    apply_legacy_label_mappings(doc, context, mapping)
    apply_table_cell_mappings(doc, context, mapping)
    apply_repeat_table_mappings(doc, context, mapping)
    apply_single_row_table_mappings(doc, context, mapping)
    apply_paragraph_mappings(doc, context, mapping)
    apply_checkbox_mappings(doc, context, mapping)
    apply_dap_default_mappings(doc, context)
    apply_dap_financing_checkboxes(doc, context)
    apply_dap_narrative_fallbacks(doc, context)

    # Ne pas nettoyer les placeholders par défaut : cela casse la structure du DAP.
    if mapping.get("cleanup_enabled") is True:
        for placeholder in mapping.get("cleanup_placeholders", []):
            replace_text_everywhere(doc, placeholder, "")

    return {
        "mapping_used": True,
        "mapping_path": str(mapping_path),
        "mapping_name": mapping.get("mapping_name"),
        "version": mapping.get("version"),
    }



# =========================================================
# DAP POST-TRAITEMENT XML SDT - V5
# =========================================================

def force_value(value, default):
    if value is None:
        return default
    if isinstance(value, str) and value.strip() == "":
        return default
    return value


def make_activity_phrase(value):
    text = str(value or "animation touristique et loisirs").strip()
    low = text.lower()
    if low.startswith("création") or low.startswith("creation"):
        return "une offre liée à la " + text
    if low.startswith("un ") or low.startswith("une ") or low.startswith("des "):
        return text
    return "une offre d’" + text


def get_installation_address(context):
    projet = context.get("projet", {}) if isinstance(context.get("projet"), dict) else {}
    entreprise = context.get("entreprise", {}) if isinstance(context.get("entreprise"), dict) else {}

    return force_value(
        projet.get("adresse_installations")
        or projet.get("adresse_site")
        or projet.get("lieu_realisation")
        or context.get("adresse_installations")
        or context.get("adresse_site")
        or entreprise.get("adresse_siege")
        or projet.get("ville_region")
        or context.get("ville_region"),
        "Agadir, Souss Massa"
    )


def get_project_city(context):
    projet = context.get("projet", {}) if isinstance(context.get("projet"), dict) else {}
    return force_value(
        projet.get("ville_region")
        or projet.get("region")
        or context.get("ville_region")
        or context.get("region"),
        "Agadir, Souss Massa"
    )


def get_dap_narrative_texts(context):
    projet = context.get("projet", {}) if isinstance(context.get("projet"), dict) else {}
    entreprise = context.get("entreprise", {}) if isinstance(context.get("entreprise"), dict) else {}

    ville = get_project_city(context)
    activite = force_value(
        projet.get("activites_envisagees") or projet.get("objet") or projet.get("description") or context.get("activite"),
        "animation touristique et loisirs"
    )
    activite_phrase = make_activity_phrase(activite)
    societe = force_value(entreprise.get("raison_sociale"), "la société")

    return {
        "market": str(force_value(
            context.get("synthese_etude_marche") or projet.get("synthese_etude_marche"),
            f"Le projet s’inscrit dans un marché touristique régional en développement, porté par la croissance de la demande pour des expériences de loisirs, d’animation et de découverte. À {ville}, l’activité touristique bénéficie d’une fréquentation nationale et internationale soutenue, ainsi que d’une demande locale pour des offres structurées, sécurisées et accessibles. L’offre proposée par {societe} répond à un besoin d’activités complémentaires aux prestations d’hébergement, de restauration et de loisirs existantes. Le positionnement du projet repose sur la qualité de service, la proximité avec les clientèles touristiques, la diversification de l’expérience proposée et la capacité à générer une fréquentation régulière. L’adéquation entre l’offre et la demande est favorable, notamment grâce au développement du tourisme expérientiel, familial et de loisirs dans la destination."
        )),
        "supply": str(force_value(
            context.get("strategie_approvisionnement") if isinstance(context.get("strategie_approvisionnement"), str) else None,
            "L’approvisionnement du projet sera organisé autour de fournisseurs locaux et nationaux sélectionnés sur la base de la qualité, de la disponibilité, du respect des délais et de la compétitivité des prix. Les principaux achats concernent les équipements d’exploitation, les consommables, la maintenance, les prestations techniques et les services de support nécessaires au bon fonctionnement de l’activité. La société prévoit de formaliser ses relations avec les fournisseurs à travers des devis, bons de commande et conditions de règlement négociées. Cette organisation permettra de sécuriser les approvisionnements, de réduire les risques de rupture et de favoriser l’intégration locale."
        )),
        "commercial": str(force_value(
            context.get("strategie_commerciale") if isinstance(context.get("strategie_commerciale"), str) else None,
            "La stratégie commerciale du projet reposera sur une approche multicanale combinant la vente directe, les réservations en ligne, les partenariats avec les agences de voyages, les établissements touristiques, les opérateurs locaux et les prescripteurs de la destination. L’offre sera adressée aux touristes nationaux, touristes internationaux, familles, groupes, entreprises et clientèle locale. La politique tarifaire sera adaptée selon la saison, le volume, les groupes et les offres packagées afin d’optimiser le taux de remplissage et le chiffre d’affaires. La communication s’appuiera sur la visibilité digitale, les réseaux sociaux, les partenariats touristiques et le bouche-à-oreille."
        )),
        "differentation": str(force_value(
            context.get("facteurs_differenciation") or projet.get("facteurs_differenciation"),
            f"Les facteurs de différenciation du projet reposent sur {activite_phrase} structurée, une expérience client encadrée, la qualité de l’accueil, la sécurité des prestations, la proximité avec les clientèles ciblées et l’adaptation de l’offre aux besoins des touristes et des familles. Le projet se distingue également par son ancrage régional, sa capacité à créer des partenariats avec les acteurs touristiques locaux et son potentiel de contribution à l’attractivité de la destination. Cette différenciation permettra de renforcer la compétitivité de l’entreprise face aux offres informelles ou peu structurées."
        )),
        "attractiveness": str(force_value(
            context.get("attractivite_touristique") or projet.get("attractivite_touristique"),
            f"Le projet contribuera à renforcer l’attractivité touristique de la destination en proposant une offre complémentaire d’animation, de loisirs et d’expérience client. Il permettra d’enrichir le parcours des visiteurs, d’augmenter les possibilités d’activités disponibles sur le territoire et d’améliorer la rétention touristique. Par son implantation à {ville}, le projet participera à la valorisation de l’écosystème local, à la création d’emplois directs et indirects et au développement d’une offre touristique plus diversifiée et mieux structurée."
        )),
    }


def replace_paragraph_text_xml(p, new_text):
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    W = "{" + W_NS + "}"

    pPr = p.find(W + "pPr")
    saved_pPr = None

    if pPr is not None:
        saved_pPr = etree.fromstring(etree.tostring(pPr))

    for child in list(p):
        p.remove(child)

    if saved_pPr is not None:
        p.append(saved_pPr)

    if str(new_text).strip() != "":
        r = etree.SubElement(p, W + "r")
        t = etree.SubElement(r, W + "t")
        t.set(XML_SPACE, "preserve")
        t.text = str(new_text)


def replace_cell_text_xml(tc, new_text):
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{" + W_NS + "}"
    ns = {"w": W_NS}
    paragraphs = tc.xpath("./w:p", namespaces=ns)
    if paragraphs:
        replace_paragraph_text_xml(paragraphs[0], new_text)
        for p in paragraphs[1:]:
            replace_paragraph_text_xml(p, "")
    else:
        p = etree.SubElement(tc, W + "p")
        replace_paragraph_text_xml(p, new_text)


def patch_dap_docx_xml_after_save(docx_path, context):
    """
    Corrige le DAP après sauvegarde Word.
    V5 ajoute :
    1) suppression des consignes résiduelles après la synthèse marché,
    2) nettoyage ciblé des placeholders XXX/Choisir/Raison sociale dans les tableaux,
    3) correction forte Adresse des installations + Ville/Région du Projet.
    """
    docx_path = Path(docx_path)
    tmp_path = docx_path.with_suffix(".patched.docx")

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
    W = "{" + W_NS + "}"
    W14 = "{" + W14_NS + "}"
    ns = {"w": W_NS, "w14": W14_NS}

    narratives = get_dap_narrative_texts(context)
    address = str(get_installation_address(context))
    city = str(get_project_city(context))

    rules = [
        ("Synthétiser l’étude de marché réalisée", narratives["market"]),
        ("Décrire brièvement les principales contraintes liées aux fournisseurs", narratives["supply"]),
        ("Décrire brièvement la stratégie commerciale", narratives["commercial"]),
        ("Décrire brièvement les facteurs de différenciation", narratives["differentation"]),
        ("Décrire succinctement l’attractivité touristique estimée", narratives["attractiveness"]),
    ]

    instruction_needles = [
        "Analyse de la demande",
        "Tendances d’évolution et taux de pénétration du marché",
        "Evolution de la taille de marché",
        "Drivers des ventes dans le marché",
        "Un marché à très forte demande",
        "Un marché très concurrencé",
        "Autres (à préciser)",
        "Analyse de l’offre",
        "Analyser l’environnement du site d’implantation",
        "Analyser l’offre future dans l’environnement immédiat",
        "Adéquation de l’offre et de la demande",
        "Analyser les principaux indicateurs d’adéquation",
        "Contribution à l’Ecosystème du tourisme",
    ]

    exact_placeholders = {
        "Choisir", "XXX%", "XX%", "XXX MMAD/an", "XXXXXXXX", "XXXXXXX", "XXXXX", "X XXX",
        "XXXX MAD", "XX personnes", "XX minutes", "XX jours", "XX jours/Cash", "Description de l’offre",
        "(Marché/clientèle cible…)", "Rationnel du choix de la stratégie du pricing",
        "Raisons sociales (si applicable)", "Ville, Pays", "e.g., Légumes", "e.g., Location des bus",
        "e.g., Location des quads", "e.g., Animation orchestre", "e.g., Animation\norchestre",
    }

    def normalized(s):
        s = str(s or "").lower().replace("\xa0", " ")
        s = s.replace("’", "'").replace("œ", "oe")
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def p_text(p):
        return "".join(t.text or "" for t in p.findall(".//" + W + "t"))

    def element_text(el):
        return "".join(t.text or "" for t in el.xpath(".//w:t", namespaces=ns))

    def bool_fin(paths):
        for path in paths:
            v = get_context_value(context, path, None)
            if v not in [None, "", [], {}]:
                return to_number(v) > 0 or str(v).strip() in ["☑", "true", "True", "1", "oui", "Oui"]
        return False

    financing = {
        "Autofinancement": bool_fin([
            "financement_expert.fonds_propres_kmad", "financement_expert.fonds_propres_mad",
            "financement_pi.fonds_propres_kmad", "projet.mode_financement.fonds_propres",
            "projet.mode_financement.autofinancement",
        ]),
        "CMT": bool_fin([
            "financement_expert.credit_bancaire_kmad", "financement_expert.credit_bancaire_mad",
            "financement_pi.credit_bancaire_kmad", "projet.mode_financement.credit_bancaire",
            "projet.mode_financement.cmt",
        ]),
        "Financement participatif": bool_fin([
            "financement_expert.financement_participatif_kmad", "financement_expert.financement_participatif_mad",
            "projet.mode_financement.financement_participatif",
        ]),
        "Crédit fournisseur": bool_fin([
            "financement_expert.credit_fournisseur_kmad", "financement_expert.credit_fournisseur_mad",
            "projet.mode_financement.credit_fournisseur",
        ]),
        "Leasing": bool_fin([
            "financement_expert.leasing_kmad", "financement_expert.leasing_mad", "projet.mode_financement.leasing",
        ]),
    }

    patched = {
        "narrative_replacements": 0,
        "instruction_cleanup": 0,
        "placeholder_cleanup": 0,
        "address_fixes": 0,
        "checkbox_replacements": 0,
    }

    with zipfile.ZipFile(docx_path, "r") as zin:
        file_data = {name: zin.read(name) for name in zin.namelist()}

    root_xml = etree.fromstring(file_data["word/document.xml"])

    # 1) Remplacement des paragraphes narratifs.
    for p in root_xml.xpath(".//w:p", namespaces=ns):
        text = p_text(p)
        norm_text = normalized(text)
        if not norm_text:
            continue
        for needle, replacement in rules:
            if normalized(needle) in norm_text:
                replace_paragraph_text_xml(p, replacement)
                patched["narrative_replacements"] += 1
                for sdt in p.xpath("ancestor::w:sdt[1]", namespaces=ns):
                    for hdr in sdt.xpath(".//w:showingPlcHdr", namespaces=ns):
                        parent = hdr.getparent()
                        if parent is not None:
                            parent.remove(hdr)
                break

    # 2) Supprimer les consignes restantes après la synthèse marché.
    for p in root_xml.xpath(".//w:p", namespaces=ns):
        text = p_text(p)
        norm_text = normalized(text)
        if any(normalized(x) in norm_text for x in instruction_needles):
            replace_paragraph_text_xml(p, "")
            patched["instruction_cleanup"] += 1

    # 3) Nettoyage ciblé des placeholders visibles dans les tableaux.
    for t in root_xml.xpath(".//w:t", namespaces=ns):
        raw = t.text or ""
        compact = re.sub(r"\s+", " ", raw).strip()
        if compact in exact_placeholders:
            t.text = ""
            patched["placeholder_cleanup"] += 1
            continue
        if re.fullmatch(r"Animation\s+[1-6]", compact):
            t.text = ""
            patched["placeholder_cleanup"] += 1
            continue

        # Ne jamais supprimer le vrai libellé Raison sociale dans la fiche signalétique.
        if compact == "Raison sociale":
            p_parent = t
            while p_parent is not None and p_parent.tag != W + "p":
                p_parent = p_parent.getparent()
            tbl_text = ""
            if p_parent is not None:
                tbls = p_parent.xpath("ancestor::w:tbl[1]", namespaces=ns)
                if tbls:
                    tbl_text = element_text(tbls[0])
            if any(key in tbl_text for key in ["PRINCIPAUX CONCURRENTS", "PRINCIPAUX FOURNISSEURS"]):
                t.text = ""
                patched["placeholder_cleanup"] += 1

    # 4) Correction forte Adresse des installations / Ville-Région du Projet.
    for tbl in root_xml.xpath(".//w:tbl", namespaces=ns):
        if "Adresse des installations" not in element_text(tbl):
            continue
        rows = tbl.xpath("./w:tr", namespaces=ns)
        for idx, tr in enumerate(rows):
            row_text = element_text(tr)
            cells = tr.xpath(".//w:tc", namespaces=ns)
            if "Adresse des installations" in row_text and len(cells) >= 2:
                replace_cell_text_xml(cells[0], "Adresse des installations")
                replace_cell_text_xml(cells[1], address)
                patched["address_fixes"] += 1
            if "du Projet" in row_text and len(cells) >= 2:
                replace_cell_text_xml(cells[0], "Ville/Région du Projet")
                replace_cell_text_xml(cells[1], city)
                patched["address_fixes"] += 1

    # 5) Correction des checkboxes visibles et de l'état Word.
    for p in root_xml.xpath(".//w:p", namespaces=ns):
        text = p_text(p)
        for label, checked in financing.items():
            if label in text:
                symbol = "☑" if checked else "☐"
                for chk in p.xpath(".//w14:checked", namespaces=ns):
                    chk.set(W14 + "val", "1" if checked else "0")
                    chk.set(W14 + "checked", "1" if checked else "0")
                changed = False
                for t in p.findall(".//" + W + "t"):
                    if t.text in ["☐", "☑", "□"]:
                        t.text = symbol
                        changed = True
                if changed:
                    patched["checkbox_replacements"] += 1

    file_data["word/document.xml"] = etree.tostring(
        root_xml,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True
    )

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in file_data.items():
            zout.writestr(name, data)

    tmp_path.replace(docx_path)
    return patched


def render_dap_docx(output_path, context):
    template_path = find_file("DAP_template.docx")

    if template_path is None:
        raise FileNotFoundError("DAP_template.docx introuvable dans /templates")

    safe_context = add_context_defaults(copy.deepcopy(context))
    tmp_docx = Path(tempfile.gettempdir()) / f"tmp_dap_{uuid.uuid4()}.docx"

    doc_tpl = DocxTemplate(str(template_path))
    doc_tpl.render(safe_context, jinja_env=DOCX_JINJA_ENV)
    doc_tpl.save(str(tmp_docx))

    doc = Document(str(tmp_docx))
    info = apply_dap_mapping_file(doc, safe_context)
    doc.save(str(output_path))

    try:
        post_info = patch_dap_docx_xml_after_save(output_path, safe_context)
        if isinstance(info, dict):
            info["xml_postprocess"] = post_info
    except Exception as e:
        if isinstance(info, dict):
            info["xml_postprocess_error"] = str(e)

    return info


# Ancien nom conservé pour compatibilité.
def render_dap_with_mapping(output_path, context):
    return render_dap_docx(output_path, context)


# =========================================================
# WORD / DOCUMENTS JURIDIQUES
# =========================================================

def apply_legal_doc_defaults(doc, context, template_name=None):
    dossier = context.get("dossier", {})
    entreprise = context.get("entreprise", {})
    dirigeant = context.get("dirigeant", {})
    banque = context.get("banque", {})

    nom = dirigeant.get("nom", "")
    qualite = dirigeant.get("qualite", dirigeant.get("fonction", "Gérant")) or "Gérant"
    cin = dirigeant.get("cin", "")
    raison = entreprise.get("raison_sociale", "")
    forme = entreprise.get("forme_juridique", "")
    rc = entreprise.get("rc", "")
    lieu = dossier.get("lieu_signature", "")
    date_sig = dossier.get("date_signature", "")

    banque_nom = banque.get("nom", banque.get("banque_partenaire", ""))
    banque_forme = banque.get("forme_juridique", banque.get("forme_juridique_banque", "Société Anonyme"))
    banque_capital = banque.get("capital", banque.get("capital_social_banque", ""))
    banque_siege = banque.get("siege", banque.get("siege_social", ""))

    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text

        if "Je soussigné(e)" in text and "agissant au nom" in text:
            paragraph.text = (
                f"Je soussigné(e) {nom} ({qualite}), "
                f"agissant au nom et pour le compte de {raison} {forme} :"
            )

        elif "Je soussigné (prénom, nom)" in text or ("Je soussigné" in text and "signataire de la convention" in text):
            paragraph.text = (
                f"Je soussigné {nom}, en sa qualité de {qualite}, "
                f"signataire de la convention d’investissement avec MAROC PME, "
                f"titulaire de la CIN N° {cin} et agissant au nom et pour le compte "
                f"de la société {raison} {forme}, inscrite au registre du commerce N° {rc} ;"
            )

        elif "Autorise ma banque partenaire" in text:
            paragraph.text = (
                f"• Autorise ma banque partenaire, {banque_nom}, {banque_forme} "
                f"au capital de {banque_capital} dirhams et dont le siège social est sis au {banque_siege}, "
                f"à transférer à Maroc PME, dans le cadre du Programme GO SIYAHA – volet soutien à l’investissement, "
                f"mon dossier de candidature, le dossier bancaire (rating, accord subordonné pour le prêt bancaire, "
                f"et le schéma de financement bancaire retenu) et une copie du contrat de crédit signé qu’elle aura constitué ;"
            )

        elif "Objet" in text and "justificatifs déposés" in text:
            paragraph.text = (
                f"Objet : Déclaration sur l’honneur relative aux justificatifs déposés par la société {raison} "
                f"pour sa candidature au programme GO SIYAHA – Volet soutien à l’investissement."
            )

        elif "Fait à" in text and ("Le" in text or "le" in text or "………" in text):
            paragraph.text = f"Fait à : {lieu}    Le : {date_sig}"


def render_legal_docx(template_name, output_path, context):
    template_path = find_file(template_name)

    if template_path is None:
        raise FileNotFoundError(f"{template_name} introuvable dans /templates")

    safe_context = add_context_defaults(copy.deepcopy(context))
    tmp_docx = Path(tempfile.gettempdir()) / f"tmp_docx_{uuid.uuid4()}.docx"

    doc_tpl = DocxTemplate(str(template_path))
    doc_tpl.render(safe_context, jinja_env=DOCX_JINJA_ENV)
    doc_tpl.save(str(tmp_docx))

    doc = Document(str(tmp_docx))
    apply_legal_doc_defaults(doc, safe_context, template_name)
    doc.save(str(output_path))


def render_docx(template_name, output_path, context):
    if template_name == "DAP_template.docx":
        render_dap_docx(output_path, context)
    else:
        render_legal_docx(template_name, output_path, context)


# =========================================================
# ROUTES DEBUG
# =========================================================

@app.get("/")
def root():
    return {
        "status": "ok",
        "message": "GO SIYAHA filler API",
        "code_version": CODE_VERSION,
        "base_dir": str(BASE_DIR),
        "templates_dir": str(TEMPLATE_DIR),
        "mappings_dir": str(MAPPING_DIR),
    }


@app.get("/health")
def health():
    return {
        "status": "ok",
        "code_version": CODE_VERSION,
        "base_dir": str(BASE_DIR),
        "templates": {
            "DAP_template.docx": find_file("DAP_template.docx") is not None,
            "BP_template.xlsx": find_file("BP_template.xlsx") is not None,
            "demande_participation.docx": find_file("demande_participation.docx") is not None,
            "engagement_capacite_financiere.docx": find_file("engagement_capacite_financiere.docx") is not None,
            "engagement_autorisations.docx": find_file("engagement_autorisations.docx") is not None,
            "declaration_honneur_justificatifs.docx": find_file("declaration_honneur_justificatifs.docx") is not None,
        },
        "mappings": {
            "mapping_bp_istitmar.json": find_file("mapping_bp_istitmar.json") is not None,
            "mapping_dap_istitmar.json": find_file("mapping_dap_istitmar.json") is not None,
        },
    }


@app.get("/debug-bp")
def debug_bp():
    template_path = find_file("BP_template.xlsx")
    mapping_path = find_file("mapping_bp_istitmar.json")

    if template_path is None or mapping_path is None:
        return JSONResponse(
            status_code=404,
            content={
                "status": "error",
                "BP_template.xlsx": template_path is not None,
                "mapping_bp_istitmar.json": mapping_path is not None,
            }
        )

    try:
        with open(mapping_path, "r", encoding="utf-8") as f:
            mapping = json.load(f)

        return {
            "status": "ok",
            "template_path": str(template_path),
            "mapping_path": str(mapping_path),
            "mapping_name": mapping.get("mapping_name"),
            "version": mapping.get("version"),
            "sections": {
                "scalar_mappings": len(mapping.get("scalar_mappings", [])),
                "financement_pi_mappings": len(mapping.get("financement_pi_mappings", [])),
                "table_mappings": len(mapping.get("table_mappings", [])),
                "cpc_mappings": len(mapping.get("cpc_mappings", [])),
                "bilan_mappings": len(mapping.get("bilan_mappings", [])),
                "impacts_mappings": len(mapping.get("impacts_mappings", [])),
            },
        }

    except Exception as e:
        return JSONResponse(status_code=500, content={"status": "error", "error": str(e)})


@app.get("/debug-mapping")
def debug_mapping():
    mapping_path = find_file("mapping_dap_istitmar.json")

    if mapping_path is None:
        return JSONResponse(
            status_code=404,
            content={"status": "error", "error": "mapping_dap_istitmar.json introuvable"}
        )

    try:
        text = mapping_path.read_text(encoding="utf-8")
        data = json.loads(text)

        return {
            "status": "ok",
            "mapping_path": str(mapping_path),
            "version": data.get("version"),
            "mapping_name": data.get("mapping_name"),
            "sections": list(data.keys()),
        }

    except Exception as e:
        lines = text.splitlines() if "text" in locals() else []

        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "mapping_path": str(mapping_path),
                "error": str(e),
                "around_line_281": lines[276:286],
            }
        )


# =========================================================
# ROUTE PRINCIPALE
# =========================================================

@app.post("/fill")
async def fill(request: Request):
    try:
        payload = await request.json()
        data = payload.get("data", payload)

        selected_template = (
            payload.get("selected_template")
            or data.get("selected_template")
            or "dossier_complet"
        )

        aliases = {
            "bp": "bp_excel",
            "dap": "dap_word",
            "demande_honneur": "declaration_honneur_justificatifs",
            "engagement_capacite": "engagement_capacite_financiere",
            "declaration_factures": "declaration_honneur_justificatifs",
        }

        selected_template = aliases.get(selected_template, selected_template)

        context = {
            **data,
            "selected_template": selected_template,
            "dossier": data.get("dossier", {}),
            "entreprise": data.get("entreprise", {}),
            "dirigeant": data.get("dirigeant", {}),
            "projet": data.get("projet", {}),
            "investissements": data.get("investissements", {}),
            "emplois": data.get("emplois", {}),
            "banque": data.get("banque", {}),
            "financement_pi": data.get("financement_pi", {}),
            "cpc_historique": data.get("cpc_historique", {}),
            "cpc_previsionnel": data.get("cpc_previsionnel", {}),
            "bilan_historique": data.get("bilan_historique", {}),
            "bilan_previsionnel": data.get("bilan_previsionnel", {}),
            "impacts_historique": data.get("impacts_historique", {}),
            "impacts_previsionnels": data.get("impacts_previsionnels", {}),
            "hypotheses_financieres": data.get("hypotheses_financieres", {}),
            "financement_expert": data.get("financement_expert", {}),
            "financement_checkbox": data.get("financement_checkbox", {}),
            "dap": data.get("dap", {}),
            "gamme_services": data.get("gamme_services", []),
            "marche_cibles": data.get("marche_cibles", []),
            "concurrents": data.get("concurrents", []),
            "fournisseurs": data.get("fournisseurs", []),
            "clients_creneaux": data.get("clients_creneaux", []),
            "politique_prix": data.get("politique_prix", []),
            "capacites_animation": data.get("capacites_animation", []),
            "emplois_directs_profils": data.get("emplois_directs_profils", []),
            "recettes_devises": data.get("recettes_devises", []),
            "achats_devises": data.get("achats_devises", {}),
            "ca_previsionnel_dap": data.get("ca_previsionnel_dap", {}),
            "synthese_etude_marche": data.get("synthese_etude_marche", ""),
            "strategie_approvisionnement": data.get("strategie_approvisionnement", {}),
            "strategie_commerciale": data.get("strategie_commerciale", {}),
            "facteurs_differenciation": data.get("facteurs_differenciation", ""),
            "attractivite_touristique": data.get("attractivite_touristique", ""),
        }

        context = add_context_defaults(context)

        dossier = context["dossier"]
        entreprise = context["entreprise"]

        identifiant = safe_name(dossier.get("identifiant", "DOSSIER"))
        societe = safe_name(entreprise.get("raison_sociale", "SOCIETE"))

        job_id = str(uuid.uuid4())
        tmp_dir = Path(tempfile.gettempdir()) / f"go_siyaha_{job_id}"
        tmp_dir.mkdir(parents=True, exist_ok=True)

        output_zip = tmp_dir / f"{identifiant}_{societe}_GO_SIYAHA.zip"
        generated_files = []
        generation_errors = []
        debug_info = {
            "selected_template": selected_template,
            "generated_files": [],
            "errors": [],
        }

        def add_bp_safe():
            try:
                output_path = tmp_dir / f"{identifiant}_{societe}_BP_GO_SIYAHA.xlsx"
                written = render_bp_excel(output_path, context)
                generated_files.append(output_path)
                debug_info["bp_written"] = written
                debug_info["generated_files"].append(output_path.name)
            except Exception as e:
                error = f"BP_template.xlsx : {str(e)}"
                generation_errors.append(error)
                debug_info["errors"].append(error)

        def add_dap_safe():
            try:
                output_path = tmp_dir / f"{identifiant}_{societe}_DAP_GO_SIYAHA.docx"
                info = render_dap_docx(output_path, context)
                generated_files.append(output_path)
                debug_info["dap_info"] = info
                debug_info["generated_files"].append(output_path.name)
            except Exception as e:
                error = f"DAP_template.docx : {str(e)}"
                generation_errors.append(error)
                debug_info["errors"].append(error)

        def add_legal_safe(template_name, output_suffix):
            try:
                output_path = tmp_dir / f"{identifiant}_{societe}_{output_suffix}.docx"
                render_legal_docx(template_name, output_path, context)
                generated_files.append(output_path)
                debug_info["generated_files"].append(output_path.name)
            except Exception as e:
                error = f"{template_name} : {str(e)}"
                generation_errors.append(error)
                debug_info["errors"].append(error)

        if selected_template == "bp_excel":
            add_bp_safe()

        elif selected_template == "dap_word":
            add_dap_safe()

        elif selected_template == "demande_participation":
            add_legal_safe("demande_participation.docx", "demande_participation")

        elif selected_template == "engagement_capacite_financiere":
            add_legal_safe("engagement_capacite_financiere.docx", "engagement_capacite_financiere")

        elif selected_template == "engagement_autorisations":
            add_legal_safe("engagement_autorisations.docx", "engagement_autorisations")

        elif selected_template == "declaration_honneur_justificatifs":
            add_legal_safe("declaration_honneur_justificatifs.docx", "declaration_honneur_justificatifs")

        elif selected_template in ["dossier_complet", "all", "tout"]:
            add_dap_safe()
            add_bp_safe()
            add_legal_safe("demande_participation.docx", "demande_participation")
            add_legal_safe("engagement_capacite_financiere.docx", "engagement_capacite_financiere")
            add_legal_safe("engagement_autorisations.docx", "engagement_autorisations")
            add_legal_safe("declaration_honneur_justificatifs.docx", "declaration_honneur_justificatifs")

        else:
            return JSONResponse(
                status_code=400,
                content={
                    "error": f"selected_template inconnu : {selected_template}",
                    "allowed": [
                        "bp_excel",
                        "dap_word",
                        "demande_participation",
                        "engagement_capacite_financiere",
                        "engagement_autorisations",
                        "declaration_honneur_justificatifs",
                        "dossier_complet",
                    ],
                },
            )

        if generation_errors:
            error_report = tmp_dir / "rapport_erreurs_generation.txt"
            error_report.write_text(
                "Erreurs de génération GO SIYAHA\n\n" + "\n".join(generation_errors),
                encoding="utf-8",
            )
            generated_files.append(error_report)

        debug_path = tmp_dir / "debug_generation.json"
        debug_path.write_text(
            json.dumps(debug_info, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        generated_files.append(debug_path)

        if not generated_files:
            return JSONResponse(status_code=400, content={"error": "Aucun fichier généré."})

        with zipfile.ZipFile(output_zip, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in generated_files:
                zipf.write(file_path, arcname=file_path.name)

        return FileResponse(
            path=str(output_zip),
            media_type="application/zip",
            filename=output_zip.name,
        )

    except FileNotFoundError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
